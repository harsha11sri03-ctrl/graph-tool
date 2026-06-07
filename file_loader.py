import pandas as pd
import os

try:
    from docx import Document
except:
    Document = None


def load_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext in [".xlsx", ".xls"]:
            return pd.read_excel(file_path)

        elif ext == ".csv":
            return pd.read_csv(file_path)

        elif ext == ".txt":
            return pd.read_csv(file_path, delim_whitespace=True)

        elif ext == ".docx" and Document:
            doc = Document(file_path)

            if not doc.tables:
                print("No tables found")
                return None

            all_tables = []

            for table in doc.tables:
                data = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                ]

                if len(data) < 2:
                    continue

                header = data[0]
                rows = data[1:]

                # Fix uneven columns
                fixed_rows = []
                for row in rows:
                    if len(row) > len(header):
                        row = row[:len(header)]
                    elif len(row) < len(header):
                        row += [""] * (len(header) - len(row))
                    fixed_rows.append(row)

                df = pd.DataFrame(fixed_rows, columns=header)
                all_tables.append(df)

            return pd.concat(all_tables, ignore_index=True, sort=False)

        else:
            print("Unsupported file")
            return None

    except Exception as e:
        print("Error:", e)
        return None